import io
import pdfplumber
import PyPDF2
import docx

class ResumeParserService:
    @classmethod
    def extract_text_from_pdf(cls, file_bytes):
        """
        Extracts text from a PDF file using pdfplumber, falling back to PyPDF2.
        """
        extracted_text = ""
        # Try pdfplumber first
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        extracted_text += text + "\n"
            if extracted_text.strip():
                return extracted_text.strip()
        except Exception as e:
            print(f"pdfplumber extraction failed: {e}. Trying PyPDF2...")
            
        # Fallback to PyPDF2
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    extracted_text += text + "\n"
            return extracted_text.strip()
        except Exception as e:
            print(f"PyPDF2 extraction failed: {e}")
            return ""

    @classmethod
    def extract_text_from_docx(cls, file_bytes):
        """
        Extracts text from a DOCX file using python-docx.
        """
        try:
            doc_file = io.BytesIO(file_bytes)
            doc = docx.Document(doc_file)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            return "\n".join(full_text).strip()
        except Exception as e:
            print(f"python-docx extraction failed: {e}")
            return ""

    @classmethod
    def extract_text(cls, file_bytes, filename):
        """
        Main entry point for raw text extraction based on filename extension.
        """
        ext = filename.split('.')[-1].lower() if '.' in filename else ''
        if ext == 'pdf':
            return cls.extract_text_from_pdf(file_bytes)
        elif ext in ['docx', 'doc']:
            return cls.extract_text_from_docx(file_bytes)
        else:
            print(f"Unsupported file type: {ext}")
            return ""
